#!/usr/bin/env python3
"""Build the bilingual student artifacts for the coding-foundations retrofit.

These files are staging assets for Google Docs and Canvas. Canvas remains the
curriculum source after the artifacts are uploaded and linked by immutable IDs.
"""

from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


INK = "000000"
MUTED = "555555"
BORDER = "DADCE0"
LIGHT = "F8F9FA"
ACCENT = "0E7C7B"
CONTENT_WIDTH_DXA = 9360


def set_font(run, name: str = "Arial", size: float = 11, bold: bool = False, color: str = INK):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_paragraph_box(paragraph, fill: str = LIGHT):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "6")
        node.set(qn("w:space"), "5")
        node.set(qn("w:color"), BORDER)


def mark_header_row(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_geometry(table, widths_dxa: list[int]):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "0")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.width = Inches(widths_dxa[index] / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            tc_w.set(qn("w:w"), str(widths_dxa[index]))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def style_document(doc: Document):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    section.header_distance = Inches(0.3)
    section.footer_distance = Inches(0.3)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.08

    for name, size, color, before, after in (
        ("Heading 1", 17, INK, 12, 5),
        ("Heading 2", 14, ACCENT, 10, 4),
        ("Heading 3", 12, MUTED, 8, 3),
    ):
        style = doc.styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = False
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        style.paragraph_format.left_indent = Inches(0)
        style.paragraph_format.right_indent = Inches(0)
        style.paragraph_format.first_line_indent = Inches(0)


def add_title(doc: Document, title: str, subtitle: str):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(title)
    set_font(run, size=24, bold=False)
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(10)
    run = paragraph.add_run(subtitle)
    set_font(run, size=10.5, color=MUTED)


def add_identity(doc: Document, labels: tuple[str, str, str]):
    table = doc.add_table(rows=1, cols=3)
    set_table_geometry(table, [3120, 3120, 3120])
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    mark_header_row(table.rows[0])
    for cell, label in zip(table.rows[0].cells, labels):
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(1)
        set_font(p.add_run(f"{label}:"), bold=True)
        p = cell.add_paragraph()
        p.paragraph_format.space_after = Pt(5)
        set_font(p.add_run("____________________"), color=MUTED)


def add_instructions(doc: Document, heading: str, lines: list[str]):
    doc.add_heading(heading, level=2)
    for line in lines:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        set_font(p.add_run(line))


def add_prompt(doc: Document, label: str, prompt: str, lines: int = 2):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    set_font(p.add_run(label + " "), bold=True, color=ACCENT)
    set_font(p.add_run(prompt))
    for _ in range(lines):
        p = doc.add_paragraph("________________________________________________________________________________")
        p.paragraph_format.space_after = Pt(3)
        set_font(p.runs[0], size=9, color=MUTED)


def add_code_box(doc: Document, title: str, code_lines: list[str]):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    set_font(p.add_run(title), bold=True, color=ACCENT)
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.left_indent = Inches(0.08)
    paragraph.paragraph_format.right_indent = Inches(0.08)
    set_paragraph_box(paragraph)
    for index, line in enumerate(code_lines):
        if index:
            paragraph.add_run("\n")
        set_font(paragraph.add_run(line), name="Courier New", size=9.5)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths_dxa: list[int], row_lines: int = 1):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths_dxa)
    mark_header_row(table.rows[0])
    for cell, value in zip(table.rows[0].cells, headers):
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_font(p.add_run(value), bold=True)
        set_cell_shading(cell, LIGHT)
    for row_values in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, row_values):
            set_font(cell.paragraphs[0].add_run(value + ("\n" * (row_lines - 1))))
    set_table_geometry(table, widths_dxa)


def add_variable_table(doc: Document, headers: list[str], rows: list[list[str]], row_lines: int = 1):
    add_table(doc, headers, rows, [2300, 1700, 2300, 3060], row_lines=row_lines)


def set_document_language(doc: Document, language: str):
    """Set both core metadata and OOXML defaults for screen readers/spellcheck."""
    doc.core_properties.language = language
    styles = doc.styles.element
    doc_defaults = styles.find(qn("w:docDefaults"))
    if doc_defaults is None:
        doc_defaults = OxmlElement("w:docDefaults")
        styles.insert(0, doc_defaults)
    rpr_default = doc_defaults.find(qn("w:rPrDefault"))
    if rpr_default is None:
        rpr_default = OxmlElement("w:rPrDefault")
        doc_defaults.append(rpr_default)
    rpr = rpr_default.find(qn("w:rPr"))
    if rpr is None:
        rpr = OxmlElement("w:rPr")
        rpr_default.append(rpr)
    lang = rpr.find(qn("w:lang"))
    if lang is None:
        lang = OxmlElement("w:lang")
        rpr.append(lang)
    lang.set(qn("w:val"), language)
    lang.set(qn("w:eastAsia"), language)
    lang.set(qn("w:bidi"), language)
    for style in doc.styles:
        if style.type != 1:
            continue
        style_rpr = style.element.get_or_add_rPr()
        style_lang = style_rpr.find(qn("w:lang"))
        if style_lang is None:
            style_lang = OxmlElement("w:lang")
            style_rpr.append(style_lang)
        style_lang.set(qn("w:val"), language)


def add_page_number(doc: Document, label: str):
    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_font(footer.add_run(label), size=8, color=MUTED)


def page_break(doc: Document):
    doc.add_page_break()


TEXT = {
    "en": {
        "language": "en-US",
        "title": "Coding Foundations Passport",
        "footer": "Coding Foundations Passport",
        "subtitle": "One planning record for blocks, games, text code, and robots.",
        "identity": ("Alias / initials", "Class", "Date started"),
        "how": "Set up this passport once",
        "how_lines": [
            "Make one copy only. Rename it ClassPeriod_AssignedAlias_Coding_Passport and place it in your course coding folder. Use only a teacher-approved alias or initials—not your full name.",
            "Open this same copy every time. Do not make a new copy for each checkpoint.",
            "Complete only the checkpoint your teacher assigns. A checkpoint may be on more than one page.",
            "When you submit, capture the completed checkpoint page or section—not the whole passport.",
            "If your copy is lost, search the exact filename first. Then ask your teacher for the recovery route and use earlier Canvas submissions to restore prior evidence.",
        ],
        "remember": "A screenshot shows that code ran. Your plan, test record, and explanation show what you understand.",
        "c1": "Checkpoint 1 - Decompose and plan before blocks",
        "c1_context": "Use with the shared Code.org route or the task your teacher names.",
        "c1_prompts": [
            ("Goal:", "What must the finished program make happen?", 2),
            ("Inputs and outputs:", "What enters the system? What should the system produce?", 2),
            ("Three subproblems:", "Break the goal into three smaller jobs.", 3),
        ],
        "pseudo_title": "First pseudocode draft",
        "pseudo_lines": ["START", "  [first precise command]", "  [next command]", "  [repeat or decision if needed]", "END"],
        "literal": "Partner literal test:",
        "literal_prompt": "Where did your partner stop because the plan was unclear or incomplete? What line will you revise?",
        "c2": "Checkpoint 2 - Find a common pattern and generalize it",
        "c2_intro": "Use this common route example even if your Hour of Code activity uses different blocks.",
        "c2_code": ["MOVE forward 3 spaces", "TURN right", "MOVE forward 3 spaces", "TURN right", "MOVE forward 3 spaces", "TURN right", "MOVE forward 3 spaces", "TURN right"],
        "c2_prompts": [
            ("Repeated group:", "Bracket the smallest useful group. Which two commands repeat?", 2),
            ("Iteration benefit:", "Why is a loop better than copying these commands four times?", 2),
        ],
        "c2_headers": ["Common variable", "Type", "Starting value", "What it controls / operation"],
        "c2_rows": [["sideLength", "number", "3", ""], ["turnDirection", "string", '"right"', ""], ["keepDrawing", "Boolean", "true", ""]],
        "c2_general_title": "Generalized pseudocode",
        "c2_general_lines": ["SET sideLength TO 3", "SET turnDirection TO right", "REPEAT 4 times", "  MOVE sideLength spaces", "  TURN turnDirection", "END REPEAT"],
        "c2_transfer": ("Transfer:", "Name one pattern from today's coding activity. How is it similar to or different from the common route?", 2),
        "c3": "Checkpoint 3 - Predict, test, and revise three times",
        "c3_headers": ["Prediction before Run", "Observed result", "One code change", "Result after change"],
        "c3_rows": [["Test 1", "", "", ""], ["Test 2", "", "", ""], ["Test 3", "", "", ""]],
        "c3_revision_headers": ["Test", "Pseudocode line before", "Revised pseudocode line", "Evidence for the revision"],
        "c3_revision_rows": [["1", "", "", ""], ["2", "", "", ""], ["3", "", "", ""]],
        "c3_claim": ("Improvement claim:", "Choose one test. When I changed ___, the program ___ because ___.", 2),
        "c4": "Checkpoint 4 - Plan and test one game feature",
        "c4_prompts": [
            ("Feature goal:", "Name one challenge, reward, rule, identity choice, feedback cue, or progression feature.", 2),
            ("Trigger and result:", "WHEN __________ happens, the program SHOULD __________.", 2),
            ("Feature pseudocode:", "Write the feature logic before changing the game.", 4),
        ],
        "c4_test_headers": ["Prediction", "Observed result", "One change", "Result after change"],
        "c4_test_rows": [["What should the feature do?", "What did it do?", "What exact block/value changed?", "What improved?"]],
        "c4_claim": ("Test claim:", "The feature changed ___ when ___. My evidence is ___.", 2),
        "c5": "Checkpoint 5 - Trace and repair the boundary bug",
        "c5_intro": "Do not repair the code until you record the prediction and first result.",
        "c5_headers": ["Code element", "Type", "Purpose", "Predicted value"],
        "c5_rows": [["missionName", "", "", ""], ["rows / columns", "", "", ""], ["priorityMode", "", "", ""], ["suppliesPlaced", "", "", ""]],
        "c5_prompts": [
            ("Prediction before Run:", "For rows = 3 and columns = 4, predict the grid and suppliesPlaced.", 2),
            ("First observed result:", "Record the grid, score, and screenshot filename before repair.", 2),
        ],
        "c5_line_headers": ["Bugged inner-loop condition", "Repaired inner-loop condition"],
        "c5_line_rows": [["column < __________", "column < __________"]],
        "c5_after": ("After repair:", "Record the corrected grid, score, and screenshot filename.", 2),
        "c5_explain": ("Nested-loop explanation:", "The outer loop solves ___ because ___. The inner loop solves ___ because ___.", 3),
        "c6": "Checkpoint 6 - Create and improve a supply-grid program",
        "c6_intro": "Plan the counters and stop values here. In JavaScript, you must author the declarations and both complete nested-loop structures yourself.",
        "c6_prompts": [
            ("Problem and user:", "Who needs this grid, and what should it help them organize?", 2),
            ("Constraints:", "Choose 4 x 5 or 2 x 6. Keep every marker on screen and keep the score accurate.", 2),
        ],
        "c6_pseudo_title": "Pseudocode before JavaScript",
        "c6_pseudo_lines": ["SET mission details", "FOR each row", "  FOR each column", "    PLACE one supply", "    UPDATE the count", "  END FOR", "END FOR", "CHECK the mission result"],
        "c6_var_headers": ["My JavaScript declaration", "Type", "Purpose", "Operation I will use"],
        "c6_var_rows": [["", "string", "", "join text with +"], ["", "number", "", "multiply / add"], ["", "Boolean", "", "AND / comparison"]],
        "c6_loop_headers": ["Outer loop plan", "Inner loop plan"],
        "c6_loop_rows": [["counter: row; stop: __________", "counter: column; stop: __________"]],
        "c6_code_evidence": "Full-code evidence: Your submitted JavaScript must show both complete nested-loop structures and their bodies.",
        "c6_continue": "Checkpoint 6 continued - Test and revise",
        "c6_prediction": ("Prediction before Run:", "Predict totalExpected, suppliesPlaced, and one x/y position.", 3),
        "c6_test_headers": ["First observed result", "Required revision", "Result after revision", "Evidence filename"],
        "c6_test_rows": [["", "", "", ""]],
        "c6_final": [
            ("Operations:", "Name one string, number, and Boolean operation visible in your code.", 3),
            ("Subproblem explanation:", "Explain how the outer and inner loops solve different jobs in the supply-grid problem.", 3),
            ("Improvement claim:", "The required revision improved the program because...", 2),
        ],
        "c7": "Checkpoint 7 - Plan the RVR mission",
        "c7_notice": "Use this page for problem solving and team planning. Use the existing RVR Mission Sheet for the sketch, robot ID, first-run result, revision, final photo, and final program evidence.",
        "c7_problem": ("Mission problem:", "What recognizable word or symbol must the robot draw? What makes it challenging?", 2),
        "c7_option_headers": ["Possible solution", "Programming mode / approach", "Benefit", "Limitation"],
        "c7_option_rows": [["Solution A", "", "", ""], ["Solution B", "", "", ""]],
        "c7_select": ("Selected solution:", "Which solution will your team use, and why does it fit the problem better?", 2),
        "c7_role_headers": ["Role", "Student", "Responsibility", "Evidence due"],
        "c7_role_rows": [["Planner", "", "", ""], ["Programmer", "", "", ""], ["Tester / evidence lead", "", "", ""]],
        "c7_timeline_headers": ["Milestone", "Expected time", "Done?", "Adjustment"],
        "c7_timeline_rows": [["Plan approved", "", "", ""], ["First run", "", "", ""], ["Revision run", "", "", ""]],
        "c7_continue": "Checkpoint 7 continued - Robot pseudocode",
        "c7_pseudo": ("Write the plan:", "Include headings, speed, duration/distance, waits, and repeated actions. Keep sketch, test, and final evidence on the existing RVR Mission Sheet.", 12),
    },
    "es": {
        "language": "es-MX",
        "title": "Pasaporte de Fundamentos de Programación",
        "footer": "Pasaporte de Fundamentos de Programación",
        "subtitle": "Un registro de planificación para bloques, juegos, código de texto y robots.",
        "identity": ("Alias / iniciales", "Clase", "Fecha de inicio"),
        "how": "Configura este pasaporte una sola vez",
        "how_lines": [
            "Crea una sola copia. Nómbrala Periodo_AliasAsignado_Pasaporte_Programación y guárdala en tu carpeta del curso. Usa solamente un alias o iniciales aprobados por tu docente, no tu nombre completo.",
            "Abre esta misma copia cada vez. No hagas una copia nueva para cada punto de control.",
            "Completa solamente el punto de control que asigne tu docente. Un punto puede ocupar más de una página.",
            "Al entregar, captura la página o sección terminada, no todo el pasaporte.",
            "Si pierdes tu copia, busca primero el nombre exacto. Luego pide la ruta de recuperación y usa tus entregas anteriores de Canvas para recuperar la evidencia.",
        ],
        "remember": "Una captura muestra que el código funcionó. Tu plan, registro de pruebas y explicación muestran lo que comprendes.",
        "c1": "Punto de control 1 - Descomponer y planear antes de usar bloques",
        "c1_context": "Usa la ruta común de Code.org o la tarea que indique tu docente.",
        "c1_prompts": [
            ("Meta:", "¿Qué debe lograr el programa terminado?", 2),
            ("Entradas y salidas:", "¿Qué entra al sistema? ¿Qué debe producir?", 2),
            ("Tres subproblemas:", "Divide la meta en tres trabajos más pequeños.", 3),
        ],
        "pseudo_title": "Primer borrador de pseudocódigo",
        "pseudo_lines": ["INICIO", "  [primer comando preciso]", "  [siguiente comando]", "  [repetición o decisión si es necesaria]", "FIN"],
        "literal": "Prueba literal con un compañero:",
        "literal_prompt": "¿Dónde se detuvo tu compañero porque el plan no era claro o estaba incompleto? ¿Qué línea revisarás?",
        "c2": "Punto de control 2 - Encontrar un patrón común y generalizarlo",
        "c2_intro": "Usa este ejemplo de ruta común aunque tu actividad de Hour of Code tenga bloques diferentes.",
        "c2_code": ["AVANZAR 3 espacios", "GIRAR a la derecha", "AVANZAR 3 espacios", "GIRAR a la derecha", "AVANZAR 3 espacios", "GIRAR a la derecha", "AVANZAR 3 espacios", "GIRAR a la derecha"],
        "c2_prompts": [
            ("Grupo repetido:", "Encierra el grupo útil más pequeño. ¿Qué dos comandos se repiten?", 2),
            ("Beneficio de la iteración:", "¿Por qué un bucle es mejor que copiar estos comandos cuatro veces?", 2),
        ],
        "c2_headers": ["Variable común", "Tipo", "Valor inicial", "Qué controla / operación"],
        "c2_rows": [["sideLength", "número", "3", ""], ["turnDirection", "texto", '"derecha"', ""], ["keepDrawing", "Booleano", "true", ""]],
        "c2_general_title": "Pseudocódigo generalizado",
        "c2_general_lines": ["FIJAR sideLength EN 3", "FIJAR turnDirection EN derecha", "REPETIR 4 veces", "  AVANZAR sideLength espacios", "  GIRAR turnDirection", "FIN REPETIR"],
        "c2_transfer": ("Transferencia:", "Nombra un patrón de la actividad de hoy. ¿En qué se parece o se diferencia de la ruta común?", 2),
        "c3": "Punto de control 3 - Predecir, probar y revisar tres veces",
        "c3_headers": ["Predicción antes de ejecutar", "Resultado observado", "Un cambio en el código", "Resultado después"],
        "c3_rows": [["Prueba 1", "", "", ""], ["Prueba 2", "", "", ""], ["Prueba 3", "", "", ""]],
        "c3_revision_headers": ["N.º", "Línea de pseudocódigo anterior", "Línea revisada", "Evidencia para la revisión"],
        "c3_revision_rows": [["1", "", "", ""], ["2", "", "", ""], ["3", "", "", ""]],
        "c3_claim": ("Afirmación de mejora:", "Elige una prueba. Cuando cambié ___, el programa ___ porque ___.", 2),
        "c4": "Punto de control 4 - Planear y probar una característica del juego",
        "c4_prompts": [
            ("Meta de la característica:", "Nombra un desafío, recompensa, regla, opción de identidad, aviso o progresión.", 2),
            ("Evento y resultado:", "CUANDO __________ ocurra, el programa DEBE __________.", 2),
            ("Pseudocódigo de la característica:", "Escribe la lógica antes de cambiar el juego.", 4),
        ],
        "c4_test_headers": ["Predicción", "Resultado observado", "Un cambio", "Resultado después"],
        "c4_test_rows": [["¿Qué debe hacer?", "¿Qué hizo?", "¿Qué bloque/valor cambió?", "¿Qué mejoró?"]],
        "c4_claim": ("Afirmación de prueba:", "La característica cambió ___ cuando ___. Mi evidencia es ___.", 2),
        "c5": "Punto de control 5 - Rastrear y reparar el error de límite",
        "c5_intro": "No repares el código hasta registrar la predicción y el primer resultado.",
        "c5_headers": ["Elemento de código", "Tipo", "Propósito", "Valor previsto"],
        "c5_rows": [["missionName", "", "", ""], ["rows / columns", "", "", ""], ["priorityMode", "", "", ""], ["suppliesPlaced", "", "", ""]],
        "c5_prompts": [
            ("Predicción antes de ejecutar:", "Con rows = 3 y columns = 4, predice la cuadrícula y suppliesPlaced.", 2),
            ("Primer resultado observado:", "Registra la cuadrícula, puntuación y nombre de la captura antes de reparar.", 2),
        ],
        "c5_line_headers": ["Condición del bucle interior con error", "Condición reparada"],
        "c5_line_rows": [["column < __________", "column < __________"]],
        "c5_after": ("Después de reparar:", "Registra la cuadrícula, puntuación y nombre de la captura corregida.", 2),
        "c5_explain": ("Explicación de bucles anidados:", "El bucle exterior resuelve ___ porque ___. El interior resuelve ___ porque ___.", 3),
        "c6": "Punto de control 6 - Crear y mejorar tu programa de suministros",
        "c6_intro": "Planea aquí los contadores y valores de parada. En JavaScript, debes escribir las declaraciones y las dos estructuras completas de bucles anidados.",
        "c6_prompts": [
            ("Problema y usuario:", "¿Quién necesita esta cuadrícula y qué le ayudará a organizar?", 2),
            ("Restricciones:", "Elige 4 x 5 o 2 x 6. Mantén cada marcador en pantalla y la puntuación correcta.", 2),
        ],
        "c6_pseudo_title": "Pseudocódigo antes de JavaScript",
        "c6_pseudo_lines": ["FIJAR detalles de la misión", "PARA cada fila", "  PARA cada columna", "    COLOCAR un suministro", "    ACTUALIZAR el conteo", "  FIN PARA", "FIN PARA", "COMPROBAR el resultado"],
        "c6_var_headers": ["Mi declaración de JavaScript", "Tipo", "Propósito", "Operación que usaré"],
        "c6_var_rows": [["", "texto", "", "unir texto con +"], ["", "número", "", "multiplicar / sumar"], ["", "Booleano", "", "AND / comparación"]],
        "c6_loop_headers": ["Plan del bucle exterior", "Plan del bucle interior"],
        "c6_loop_rows": [["contador: row; parada: __________", "contador: column; parada: __________"]],
        "c6_code_evidence": "Evidencia del código: El JavaScript entregado debe mostrar las dos estructuras completas de bucles anidados y sus cuerpos.",
        "c6_continue": "Punto de control 6 (continuación) — Probar y revisar",
        "c6_prediction": ("Predicción antes de ejecutar:", "Predice totalExpected, suppliesPlaced y una posición x/y.", 3),
        "c6_test_headers": ["Primer resultado observado", "Revisión obligatoria", "Resultado después", "Nombre de evidencia"],
        "c6_test_rows": [["", "", "", ""]],
        "c6_final": [
            ("Operaciones:", "Nombra una operación de texto, número y Booleano visible en tu código.", 3),
            ("Explicación de subproblemas:", "Explica cómo los bucles exterior e interior resuelven trabajos diferentes.", 3),
            ("Afirmación de mejora:", "La revisión obligatoria mejoró el programa porque...", 2),
        ],
        "c7": "Punto de control 7 - Planear la misión de RVR",
        "c7_notice": "Usa esta página para resolver el problema y planear en equipo. Usa la Hoja de Misión de RVR existente para el boceto, ID del robot, primera prueba, revisión, foto final y programa final.",
        "c7_problem": ("Problema de la misión:", "¿Qué palabra o símbolo reconocible debe dibujar el robot? ¿Qué lo hace difícil?", 2),
        "c7_option_headers": ["Solución posible", "Modo / estrategia", "Beneficio", "Limitación"],
        "c7_option_rows": [["Solución A", "", "", ""], ["Solución B", "", "", ""]],
        "c7_select": ("Solución elegida:", "¿Cuál usará tu equipo y por qué se ajusta mejor al problema?", 2),
        "c7_role_headers": ["Rol", "Estudiante", "Responsabilidad", "Evidencia"],
        "c7_role_rows": [["Planificador", "", "", ""], ["Programador", "", "", ""], ["Pruebas / evidencia", "", "", ""]],
        "c7_timeline_headers": ["Meta intermedia", "Tiempo previsto", "¿Listo?", "Ajuste"],
        "c7_timeline_rows": [["Plan aprobado", "", "", ""], ["Primera prueba", "", "", ""], ["Prueba revisada", "", "", ""]],
        "c7_continue": "Punto de control 7 (continuación) - Pseudocódigo del robot",
        "c7_pseudo": ("Escribe el plan:", "Incluye dirección, velocidad, duración/distancia, pausas y repeticiones. Guarda el boceto, pruebas y evidencia final en la Hoja de Misión de RVR existente.", 12),
    },
}


def add_intro_line(doc: Document, value: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    set_font(p.add_run(value), bold=True, color=MUTED)


def build(lang: str, output: Path):
    t = TEXT[lang]
    doc = Document()
    style_document(doc)
    set_document_language(doc, t["language"])
    add_title(doc, t["title"], t["subtitle"])
    add_identity(doc, t["identity"])
    add_instructions(doc, t["how"], t["how_lines"])
    add_code_box(doc, "Remember" if lang == "en" else "Recuerda", [t["remember"]])

    page_break(doc)
    doc.add_heading(t["c1"], level=1)
    add_intro_line(doc, t["c1_context"])
    for label, prompt, lines in t["c1_prompts"]:
        add_prompt(doc, label, prompt, lines)
    add_code_box(doc, t["pseudo_title"], t["pseudo_lines"])
    add_prompt(doc, t["literal"], t["literal_prompt"], 2)

    page_break(doc)
    doc.add_heading(t["c2"], level=1)
    add_intro_line(doc, t["c2_intro"])
    add_code_box(doc, "Common route" if lang == "en" else "Ruta común", t["c2_code"])
    for label, prompt, lines in t["c2_prompts"]:
        add_prompt(doc, label, prompt, lines)
    add_variable_table(doc, t["c2_headers"], t["c2_rows"])
    add_code_box(doc, t["c2_general_title"], t["c2_general_lines"])
    add_prompt(doc, *t["c2_transfer"])

    page_break(doc)
    doc.add_heading(t["c3"], level=1)
    add_variable_table(doc, t["c3_headers"], t["c3_rows"], row_lines=3)
    doc.add_heading("Revise the matching pseudocode after each test" if lang == "en" else "Revisa el pseudocódigo correspondiente después de cada prueba", level=2)
    add_table(doc, t["c3_revision_headers"], t["c3_revision_rows"], [700, 2700, 2700, 3260], row_lines=2)
    add_prompt(doc, *t["c3_claim"])

    page_break(doc)
    doc.add_heading(t["c4"], level=1)
    for label, prompt, lines in t["c4_prompts"]:
        add_prompt(doc, label, prompt, lines)
    doc.add_heading("Test the planned feature" if lang == "en" else "Prueba la característica planeada", level=2)
    add_variable_table(doc, t["c4_test_headers"], t["c4_test_rows"], row_lines=3)
    add_prompt(doc, *t["c4_claim"])

    page_break(doc)
    doc.add_heading(t["c5"], level=1)
    add_intro_line(doc, t["c5_intro"])
    add_variable_table(doc, t["c5_headers"], t["c5_rows"])
    for label, prompt, lines in t["c5_prompts"]:
        add_prompt(doc, label, prompt, lines)
    add_table(doc, t["c5_line_headers"], t["c5_line_rows"], [4680, 4680], row_lines=2)
    add_prompt(doc, *t["c5_after"])
    add_prompt(doc, *t["c5_explain"])

    page_break(doc)
    doc.add_heading(t["c6"], level=1)
    add_intro_line(doc, t["c6_intro"])
    for label, prompt, lines in t["c6_prompts"]:
        add_prompt(doc, label, prompt, lines)
    add_code_box(doc, t["c6_pseudo_title"], t["c6_pseudo_lines"])
    doc.add_heading("Student-authored variables and operations" if lang == "en" else "Variables y operaciones escritas por el estudiante", level=2)
    add_variable_table(doc, t["c6_var_headers"], t["c6_var_rows"], row_lines=2)
    add_table(doc, t["c6_loop_headers"], t["c6_loop_rows"], [4680, 4680], row_lines=2)
    add_intro_line(doc, t["c6_code_evidence"])

    page_break(doc)
    doc.add_heading(t["c6_continue"], level=1)
    add_prompt(doc, *t["c6_prediction"])
    add_variable_table(doc, t["c6_test_headers"], t["c6_test_rows"], row_lines=5)
    for label, prompt, lines in t["c6_final"]:
        add_prompt(doc, label, prompt, lines)

    page_break(doc)
    doc.add_heading(t["c7"], level=1)
    add_code_box(doc, "Evidence boundary" if lang == "en" else "Límite de evidencia", [t["c7_notice"]])
    add_prompt(doc, *t["c7_problem"])
    add_variable_table(doc, t["c7_option_headers"], t["c7_option_rows"], row_lines=2)
    add_prompt(doc, *t["c7_select"])
    doc.add_heading("Team roles and timeline" if lang == "en" else "Roles y línea de tiempo", level=2)
    add_variable_table(doc, t["c7_role_headers"], t["c7_role_rows"])
    add_variable_table(doc, t["c7_timeline_headers"], t["c7_timeline_rows"])

    page_break(doc)
    doc.add_heading(t["c7_continue"], level=1)
    add_prompt(doc, *t["c7_pseudo"])

    add_page_number(doc, t["footer"])
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--output-dir", type=Path, required=True)
    args = parser.parse_args()
    build("en", args.output_dir / "Coding_Foundations_Passport_EN.docx")
    build("es", args.output_dir / "Coding_Foundations_Passport_ES.docx")


if __name__ == "__main__":
    main()
